from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth import get_user_model
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.models import User
from django.views.decorators.http import require_POST
from django.urls import reverse
from django.utils import timezone
from django.db import models
from django import forms
import io
import base64
import re
import fitz
from django.core.mail import send_mail
from django.db.models import Q
from django.conf import settings
from io import BytesIO
from docx import Document as DocxDocument
from rest_framework import generics, status
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.views import APIView
from users.models import UserProfile
from .utils import generate_document_from_template

# --- WeasyPrint Imports ---
from weasyprint import HTML, CSS
from django.template.loader import render_to_string

# We need to import the admin test function from our 'users' app

# --- Model Imports ---
from .models import (
    Case, CaseAssignment, Document, DocumentLog,
    CaseWorkflow, CaseStage, Template, 
    SignatureRequest, CaseStageLog, Meeting, DocumentDueDate, ContractTemplate, ConsultationRequest,   
)
from django.db.models import Sum, Count, Avg, F, Count
from users.models import Role
from .serializers import ContractTemplateSerializer

# --- Form Imports ---
from .forms import (
    CaseCreateForm, DocumentUploadForm,
    WorkflowCreateForm, StageCreateForm, 
    TemplateUploadForm, MeetingForm, ConsultationScheduleForm, ConsultationForm,
)

from users.views import is_admin
from .decorators import user_is_assigned_to_case

# --- Django's File handling utilities ---
from django.core.files.base import ContentFile
from django.conf import settings # <-- Import settings
from django.http import HttpResponse, JsonResponse, FileResponse
from django.views.decorators.csrf import csrf_exempt
from datetime import datetime
import os
import json

# --- HELPER FUNCTION FOR DOCX ---
def docx_find_and_replace(doc, context):
    """
    Finds and replaces text in a .docx file, preserving formatting.
    Placeholders must be in the format {{key_name}}.
    """
    # Combine paragraphs from the main body and all table cells
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)
    
    for p in all_paragraphs:
        for key, value in context.items():
            placeholder = f"{{{{{key}}}}}" # Creates {{key_name}}
            
            # We must iterate over 'runs' to preserve formatting (bold, etc.)
            # A simple p.text.replace() would wipe all formatting.
            if placeholder in p.text:
                for run in p.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
    return doc

# --- View 1: Case List (Admin) ---

@login_required
def case_dashboard_view(request):
    # 1. Determine User Roles
    is_admin_user = request.user.roles.filter(name='Admin').exists()
    is_attorney = request.user.roles.filter(name='Attorney').exists()
    
    # 2. Security Check
    if not (is_admin_user or is_attorney):
        return redirect('users:dashboard')

    # 3. Get Cases (Active only for dashboard)
    cases_list = Case.objects.filter(is_archived=False).select_related('current_stage').prefetch_related('assignments__user').order_by('-date_filed')
    
    # Filter cases: Admins see all, Attorneys see assigned only
    if not is_admin_user:
        cases_list = cases_list.filter(assignments__user=request.user)

    # --- CHART DATA CALCULATION ---
    # Group active cases by their stage name and count them
    stage_data = cases_list.values('current_stage__name').annotate(count=Count('id'))
    
    # Prepare lists for Chart.js (Handle cases with no stage as 'New')
    stage_labels = [item['current_stage__name'] if item['current_stage__name'] else 'New' for item in stage_data]
    stage_counts = [item['count'] for item in stage_data]
    # ------------------------------

    active_count = cases_list.count()
    pending_count = SignatureRequest.objects.filter(status='pending').count() if 'SignatureRequest' in globals() else 0

    # 4. Fetch Consultation Requests
    consultations = []
    
    if is_attorney:
        consultations = ConsultationRequest.objects.filter(attorney=request.user, status='Pending').order_by('-created_at')
    elif is_admin_user:
        consultations = ConsultationRequest.objects.filter(status='Pending').order_by('-created_at')

    # --- NEW: Fetch Recent Activity ---
    # Get the last 5 stage changes across all cases
    recent_activity = CaseStageLog.objects.all().select_related('case', 'stage').order_by('-timestamp_entered')[:5]

    context = {
        'cases': cases_list,
        'active_case_count': active_count,
        'pending_signature_count': pending_count,
        'is_admin_user': is_admin_user,
        'consultations': consultations,
        'stage_labels': json.dumps(stage_labels), 
        'stage_counts': json.dumps(stage_counts),
        'recent_activity': recent_activity, # <--- PASS TO TEMPLATE
    }
    
    # Ensure this points to your new template file
    return render(request, 'cases/admin_dashboard.html', context)

# --- View 2: Case Create (Admin) ---

@login_required
#@user_passes_test(is_admin) //we want attorneys to also have access to this
def case_create_view(request):
    # --- Permission Check: Only Admins and Attorneys can create cases ---
    is_admin = request.user.roles.filter(name='Admin').exists()
    is_attorney = request.user.roles.filter(name='Attorney').exists()
    
    if not (is_admin or is_attorney):
        messages.error(request, "Only admins and attorneys can create cases.")
        return redirect('users:dashboard')
    
    if request.method == 'POST':
        form = CaseCreateForm(request.POST)
        if form.is_valid():
            case = form.save() # Saves the case object

            # If a workflow was assigned, set the current_stage
            # to the first stage in that workflow (order=1).
            if case.workflow:
                first_stage = case.workflow.stages.filter(order=1).first()
                if first_stage:
                    case.current_stage = first_stage
                    case.save() # Save the change to current_stage

                    # Log the start of the workflow
                    CaseStageLog.objects.create(case=case, stage=first_stage)

            # Step 2: Get the attorney and client from the form's data
            attorney = form.cleaned_data['attorney']
            client = form.cleaned_data['client']
            
            # 💡 FIX: Create the CaseAssignment links, explicitly defining the role.
            # This is essential because the CaseAssignment model likely requires the 'role' field.
            
            # 1. Assign the Attorney
            CaseAssignment.objects.create(case=case, user=attorney)
            
            # 2. Assign the Client
            CaseAssignment.objects.create(case=case, user=client)
            
            messages.success(request, f"Successfully created case: {case.case_title}")
            # Redirect based on user role
            if is_admin:
                return redirect('cases:case-dashboard')  # Admins go to case dashboard
            else:
                return redirect('users:dashboard')  # Attorneys go to their dashboard
    else:
        # If it's a GET request, just show a blank form
        form = CaseCreateForm()
        
    context = {
        'form': form
    }
    return render(request, 'cases/case_create.html', context)

@login_required
@user_passes_test(is_admin) # Only admins can hit this URL
def case_delete_view(request, pk):
    case = get_object_or_404(Case, pk=pk)
    
    if request.method == 'POST':
        # Archive it instead of hard deleting? 
        # If you want hard delete: case.delete()
        # Here we assume hard delete based on your request:
        case.delete()
        messages.success(request, "Case deleted successfully.")
        return redirect('cases:case-dashboard')
        
    # If GET, redirect back (safety)
    return redirect('cases:case-dashboard')

# --- View 3: Case Detail (SECURED) ---
@login_required
@user_is_assigned_to_case
def case_detail_view(request, pk):
    case = get_object_or_404(Case, pk=pk)
    next_stage = None
    
    if request.method == 'POST':
        
        # --- Check which form was submitted ---
        
        # Check if the 'upload_document' button was pressed
        if 'upload_document' in request.POST:
            upload_form = DocumentUploadForm(request.POST, request.FILES)
            if upload_form.is_valid():
                doc = upload_form.save(commit=False) 
                doc.case = case
                doc.uploaded_by = request.user
                doc.save()
                DocumentLog.objects.create(document=doc, user=request.user, action="Uploaded")
                messages.success(request, f"Document '{doc.title}' uploaded successfully.")
                return redirect('cases:case-detail', pk=case.pk)
        
    
    # --- GET Request Logic ---
    documents = case.documents.all().order_by('-id')
    assignments = case.assignments.all()
    upload_form = DocumentUploadForm()
    
    # Get all stages for the case's workflow
    all_stages = None
    if case.workflow and case.current_stage: # Check both exist
        all_stages = case.workflow.stages.all()
        # --- Calculate next_stage HERE ---
        current_order = case.current_stage.order
        next_stage = all_stages.filter(order=current_order + 1).first()


    context = {
        'case': case,
        'documents': documents,
        'assignments': assignments,
        'upload_form': upload_form,
        'all_stages': all_stages,
        'next_stage': next_stage,
    }
    return render(request, 'cases/case_detail.html', context)

# --- View 4: Document Audit Log (SECURED) ---
@login_required
def document_view_and_log(request, doc_pk):
    doc = get_object_or_404(Document, pk=doc_pk)
    case = doc.case # Get the case this document belongs to

    # --- Security Check ---
    is_admin = request.user.roles.filter(name='Admin').exists()
    is_assigned = CaseAssignment.objects.filter(case=case, user=request.user).exists()
    
    if not (is_admin or is_assigned):
        # If not an admin and not assigned, forbid access
        messages.error(request, "You do not have permission to view that document.")
        return redirect('users:dashboard') # Send to their dashboard
    
    # --- If security passes, log the action ---
    DocumentLog.objects.create(
        document=doc,
        user=request.user,
        action="Viewed"
    )
    
    # And redirect to the file
    return redirect(doc.file_upload.url)

# --- View 5: Workflow List (Admin) ---
@login_required
#@user_passes_test(is_admin) //we also want attorneys to access this
def workflow_list_view(request):
    # --- Permission Check: Only Admins and Attorneys can manage workflows ---
    is_admin = request.user.roles.filter(name='Admin').exists()
    is_attorney = request.user.roles.filter(name='Attorney').exists()
    
    if not (is_admin or is_attorney):
        messages.error(request, "Only admins and attorneys can manage workflows.")
        return redirect('users:dashboard')

    workflows = CaseWorkflow.objects.all()
    context = {
        'workflows': workflows
    }
    return render(request, 'cases/workflow_list.html', context)

# --- View 6: Workflow Create (Admin) ---
@login_required
@user_passes_test(is_admin)
def workflow_create_view(request):
    if request.method == 'POST':
        form = WorkflowCreateForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "New workflow created successfully.")
            return redirect('cases:workflow-list')
    else:
        form = WorkflowCreateForm()
    
    context = {
        'form': form
    }
    return render(request, 'cases/workflow_create.html', context)

# --- View 7: Workflow Detail (Admin) ---
@login_required
@user_passes_test(is_admin)
def workflow_detail_view(request, pk):
    # Get the parent workflow
    workflow = get_object_or_404(CaseWorkflow, pk=pk)
    
    # Get all stages for this workflow, which are already ordered
    # thanks to the 'ordering' Meta option in the model
    stages = workflow.stages.all()
    
    context = {
        'workflow': workflow,
        'stages': stages
    }
    return render(request, 'cases/workflow_detail.html', context)

# --- View 8: Stage Create (Admin) ---
@login_required
@user_passes_test(is_admin)
def stage_create_view(request, workflow_pk):
    # Get the parent workflow this stage will belong to
    workflow = get_object_or_404(CaseWorkflow, pk=workflow_pk)
    
    if request.method == 'POST':
        form = StageCreateForm(request.POST)
        if form.is_valid():
            # Create the stage in memory
            stage = form.save(commit=False)
            # Assign it to the correct parent workflow
            stage.workflow = workflow
            stage.save()
            
            messages.success(request, f"New stage '{stage.name}' added to workflow.")
            # Redirect back to the detail page for that workflow
            return redirect('cases:workflow-detail', pk=workflow.pk)
    else:
        form = StageCreateForm()
        
    context = {
        'form': form,
        'workflow': workflow
    }
    return render(request, 'cases/stage_create.html', context)

# --- View 9: Advance Case Stage (NEW) ---
@login_required
@user_is_assigned_to_case
def advance_stage_view(request, case_pk):
    case = get_object_or_404(Case, pk=case_pk)
    
    # --- Security: Only Attorneys or Admins can advance a stage ---
    is_admin = request.user.roles.filter(name='Admin').exists()
    is_attorney = request.user.roles.filter(name='Attorney').exists()
    
    if not (is_admin or is_attorney):
        messages.error(request, "Only attorneys can advance a case stage.")
        return redirect('cases:case-detail', pk=case.pk)

    # --- Main Logic ---
    if not case.workflow or not case.current_stage:
        messages.error(request, "This case has no workflow to advance.")
        return redirect('cases:case-detail', pk=case.pk)

    current_order = case.current_stage.order
    
    # Find the next stage in the sequence
    next_stage = case.workflow.stages.filter(order=current_order + 1).first()
    
    if next_stage:
        now = timezone.now()
        
        # 1. Find the current, active log entry (it has no 'completed' timestamp)
        current_log_entry = CaseStageLog.objects.filter(
            case=case, 
            stage=case.current_stage,  # <--- FIX: Access it from the case object
            timestamp_completed__isnull=True
        ).first()
        
        # 2. Mark the current log entry as completed
        if current_log_entry:
            current_log_entry.timestamp_completed = now
            current_log_entry.save()
            
        # 3. Update the case's current stage
        case.current_stage = next_stage
        case.save()
        
        # 4. Create a NEW log entry for the stage we are entering
        CaseStageLog.objects.create(case=case, stage=next_stage, timestamp_entered=now)
        
        messages.success(request, f"Case stage advanced to: {next_stage.name}")
    else:
        # This was the last stage
        messages.info(request, "This case is already in its final stage.")
        
    # Redirect back to the detail page
    return redirect('cases:case-detail', pk=case.pk)

# --- View 10: Template List (Admin) ---
@login_required
def template_generation_view(request):
    templates = Template.objects.all().order_by('name')

    # Discover any contract templates placed under the project's
    # `templates/cases/` directory that end with `_contract.html`.
    contract_files = []
    templates_dir = os.path.join(settings.BASE_DIR, 'templates', 'cases')
    if os.path.isdir(templates_dir):
        for fn in sorted(os.listdir(templates_dir)):
            # include both English and Spanish variants (e.g. _contract.html and _contract_es.html)
            if fn.endswith('_contract.html') or fn.endswith('_contract_es.html'):
                contract_files.append(fn)

    # Build prettier labels server-side to avoid using template filters
    contract_choices = []
    for fn in contract_files:
        # remove suffix and replace underscores with spaces
        label = fn.replace('_contract_es.html', '').replace('_contract.html', '').replace('_', ' ')
        # Format camelCase words by inserting spaces before uppercase letters
        # and also insert space before "and" when preceded by lowercase letter
        label = re.sub(r'([a-z])([A-Z])', r'\1 \2', label)
        label = re.sub(r'([a-z])(and)', r'\1 \2', label)
        label = label.title()
        contract_choices.append((fn, label))

    context = {
        'templates': templates,
        # a list like ["vehicle_PurchaseandSale_contract.html", ...]
        'contract_files': contract_files,
        # choices list of (filename, pretty label) for template rendering
        'contract_choices': contract_choices,
        # json-encoded version to inject safely into JS
        'contract_files_json': json.dumps(contract_files),
        # lightweight placeholder map for client-side autofill (no case context here)
        'placeholder_map_json': json.dumps({})
    }
    return render(request, 'cases/template_generation.html', context)


@login_required
def contract_template_view(request):
    """Return the raw HTML of a contract template located under
    `templates/cases/` given a `file` GET parameter. This endpoint
    validates the filename to avoid path traversal and ensures only
    files ending with `_contract.html` are returned.
    """
    file_name = request.GET.get('file')
    if not file_name or '/' in file_name or '..' in file_name:
        return JsonResponse({'error': 'Invalid file parameter.'}, status=400)

    # Allow both English and Spanish contract filenames. Accepted patterns:
    # - something_contract.html
    # - something_contract_es.html
    import re
    if not re.search(r'_contract(_es)?\.html$', file_name, re.IGNORECASE):
        return JsonResponse({'error': 'Not a contract template.'}, status=400)

    templates_dir = os.path.join(settings.BASE_DIR, 'templates', 'cases')
    file_path = os.path.join(templates_dir, file_name)
    # Ensure the resolved path is actually inside the templates_dir
    try:
        real_templates_dir = os.path.realpath(templates_dir)
        real_file_path = os.path.realpath(file_path)
    except Exception:
        return JsonResponse({'error': 'Invalid path resolution.'}, status=400)

    if not real_file_path.startswith(real_templates_dir):
        return JsonResponse({'error': 'Invalid file path.'}, status=400)

    if not os.path.exists(real_file_path):
        return JsonResponse({'error': 'File not found.'}, status=404)

    try:
        with open(real_file_path, 'r', encoding='utf-8') as fh:
            content = fh.read()
    except Exception:
        return JsonResponse({'error': 'Unable to read file.'}, status=500)

    return HttpResponse(content, content_type='text/html')

# --- View 11: Template Upload (Admin) ---
@login_required
@user_passes_test(is_admin)
def template_upload_view(request):
    if request.method == 'POST':
        form = TemplateUploadForm(request.POST, request.FILES)
        if form.is_valid():
            # 1. Get the file object
            # Note: check your forms.py to see if the field is named 'file', 'template_file', etc.
            # Assuming it is 'file' based on standard Django forms:
            uploaded_file = request.FILES.get('file') 
            
            # 2. Check if it is a PDF
            if uploaded_file and uploaded_file.name.lower().endswith('.pdf'):
                try:
                    # --- THE CONVERSION MAGIC ---
                    # Read the PDF stream from memory
                    pdf_data = uploaded_file.read()
                    
                    # Open with PyMuPDF
                    with fitz.open(stream=pdf_data, filetype="pdf") as doc:
                        html_content = ""
                        # Loop through every page and convert to HTML
                        for page in doc:
                            html_content += page.get_text("html")
                    
                    # 3. Save to your Database (ContractTemplate)
                    # This ensures it shows up in your "My Custom Templates" dropdown
                    ContractTemplate.objects.create(
                        name=form.cleaned_data.get('name') or uploaded_file.name,
                        content=html_content,
                        created_by=request.user,
                        is_public=False  # Keep private by default
                    )
                    
                    messages.success(request, "PDF converted to HTML and saved as a new template!")
                    return redirect('cases:case-dashboard') # Or wherever you want to go

                except Exception as e:
                    messages.error(request, f"Error converting PDF: {e}")
                    return redirect('cases:template-upload')

            else:
                # Fallback: If it's not a PDF (e.g., standard file upload), run your old logic
                template = form.save(commit=False)
                template.uploaded_by = request.user
                template.save()
                messages.success(request, f"Template '{template.name}' uploaded successfully.")
                return redirect('cases:template-list')

    else:
        form = TemplateUploadForm()
    
    context = {
        'form': form
    }
    return render(request, 'cases/template_upload.html', context)


@login_required
@require_POST
def convert_pdf_api_view(request):
    """
    Smart Analyzer v5 (Native Scale):
    - REMOVED ZOOM: Analyzes PDF at 100% (1:1) scale to match frontend perfectly.
    - Keeps the improved 'Text Overlap' and 'Margin' filters.
    """
    uploaded_file = request.FILES.get('pdf_file')
    
    if not uploaded_file or not uploaded_file.name.lower().endswith('.pdf'):
        return JsonResponse({'error': 'Invalid file'}, status=400)

    try:
        pdf_data = uploaded_file.read()
        pages_data = []
        
        with fitz.open(stream=pdf_data, filetype="pdf") as doc:
            for page_num, page in enumerate(doc):
                # 1. RENDER PAGE AT NATIVE 1x SCALE (No Zoom)
                # This ensures coordinates match the PDF's internal 'Point' system exactly
                pix = page.get_pixmap() 
                img_data = base64.b64encode(pix.tobytes("png")).decode("utf-8")
                
                # These dimensions are now the standard PDF points (e.g. 612x792 for Letter)
                page_w = page.rect.width
                page_h = page.rect.height
                
                # SETTINGS
                MARGIN_X_PCT = 0.10   
                TEXT_Y_OFFSET = 0     # With 1x scale, we likely don't need a manual offset anymore
                
                fields = []

                def is_valid_field(rect):
                    # Margin Check
                    if rect.x0 < (page_w * MARGIN_X_PCT): return False
                    if rect.x1 > (page_w * (1 - MARGIN_X_PCT)): return False
                    
                    # Text Overlap Check
                    check_rect = fitz.Rect(rect.x0, rect.y0 - 2, rect.x1, rect.y0)
                    text_inside = page.get_text("text", clip=check_rect).strip()
                    if len(text_inside) > 0: 
                        return False 
                    return True

                def add_field(rect, f_type):
                    for existing in fields:
                        if rect.intersects(existing['rect']): return

                    # Label Guessing
                    search_rect = fitz.Rect(rect.x0 - 150, rect.y0 - 10, rect.x0, rect.y1 + 5)
                    text_nearby = page.get_text("text", clip=search_rect).strip()
                    label = text_nearby.replace(':', '').replace('\n', ' ').strip()
                    
                    fields.append({
                        # Coordinates are now purely 1:1 with the rendered image
                        'left': (rect.x0 / page_w) * 100,
                        'top': (rect.y0 / page_h) * 100,
                        'width': (rect.width / page_w) * 100,
                        'height': (rect.height / page_h) * 100,
                        'type': f_type,
                        'label': label,
                        'rect': rect
                    })

                # A. Detect Drawings
                for shape in page.get_drawings():
                    rect = shape['rect']
                    
                    # TEXT INPUTS (Horizontal Lines)
                    # Note: width > 25 is still safe at 1x scale
                    if rect.width > 25 and rect.height < 5: 
                        # Box sits ON the line
                        input_rect = fitz.Rect(rect.x0, rect.y0 - 12, rect.x1, rect.y0)
                        if is_valid_field(input_rect):
                            add_field(input_rect, 'text')
                    
                    # CHECKBOXES (Squares)
                    # Note: Adjusted min size slightly for 1x scale
                    elif 8 < rect.width < 25 and 8 < rect.height < 25:
                        ratio = rect.width / rect.height
                        if 0.8 < ratio < 1.2:
                             if is_valid_field(rect):
                                add_field(rect, 'checkbox')

                # B. Legacy "___"
                for rect in page.search_for("___"):
                    add_field(rect, 'text')
                
                # C. Checkbox Characters
                for char in ["☐", "☑", "☒"]:
                    for rect in page.search_for(char):
                        add_field(rect, 'checkbox')

                for f in fields:
                    if 'rect' in f: del f['rect']

                pages_data.append({
                    'image': img_data,
                    'width': page_w,
                    'height': page_h,
                    'fields': fields
                })
        
        return JsonResponse({'pages': pages_data})
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
# --- STEP 19: DOCUMENT GENERATION VIEW (UPGRADED)
# ---
@login_required
@user_is_assigned_to_case
def generate_document_view(request, case_pk):
    case = get_object_or_404(Case, pk=case_pk)

    #Moving out of test block to test error
    client_assignment = case.assignments.filter(user__roles__name='Client').first()
    attorney_assignment = case.assignments.filter(user__roles__name='Attorney').first()

    templates = Template.objects.filter(
        models.Q(is_public=True) | models.Q(uploaded_by=request.user)
    ).distinct()

    if request.method == 'POST':
        template_id = request.POST.get('template_id')
        if not template_id:
            messages.error(request, "Please select a template.")
        else:
            template = get_object_or_404(Template, pk=template_id)

            # --- 1. Build the Context ---
            # This matches placeholders to real data
            context = {}

            # Find the client
            if client_assignment:
                context['client_name'] = client_assignment.user.get_full_name()
                context['client_email'] = client_assignment.user.email

            # Find the attorney
            if attorney_assignment:
                context['attorney_name'] = attorney_assignment.user.get_full_name()

            context['case_title'] = case.case_title
            context['case_description'] = case.description

            # You can add any other fields from your Template.context_fields here
            # For example, if you stored {"client_address": "..."} in the case model.

            # --- 2. Generate the .docx Document ---
            try:
                # Open the template file (.docx)
                doc = DocxDocument(template.template_file.open())

                # Run our find-and-replace function
                doc = docx_find_and_replace(doc, context)

                # --- 3. Save the new file in memory ---
                file_buffer = io.BytesIO()
                doc.save(file_buffer)

                # Create a new file name
                file_name = f"Generated_{template.name.replace(' ', '_')}.docx"

                # Get the file content from the buffer
                file_content = file_buffer.getvalue()

                # --- 4. Save the new Document to the Case ---
                new_doc = Document.objects.create(
                    case=case,
                    title=f"Generated: {template.name}",
                    uploaded_by=request.user,
                    # Create a Django ContentFile from the buffer
                    file_upload=ContentFile(file_content, name=file_name)
                )

                # 5. Log the creation
                DocumentLog.objects.create(document=new_doc, user=request.user, action="Generated")

                messages.success(request, f"Document '{new_doc.title}' generated successfully.")
                return redirect('cases:case-detail', pk=case.pk)

            except Exception as e:
                print(f"!!! Error generating document: {e}")
                messages.error(request, f"Error generating document. Is the template file a valid .docx? Error: {e}")

# 1. FETCH DATABASE TEMPLATES (Add this block)
    db_templates = ContractTemplate.objects.filter(
        models.Q(is_public=True) | models.Q(created_by=request.user)
    ).order_by('name')

    # Discover any HTML contract templates placed under the project's
    # `templates/cases/` directory that end with `_contract.html`.
    templates_dir = os.path.join(settings.BASE_DIR, 'templates', 'cases')
    contract_files = []
    contract_choices = []
    if os.path.isdir(templates_dir):
        for fn in sorted(os.listdir(templates_dir)):
            # include both English and Spanish variants (e.g. _contract.html and _contract_es.html)
            if fn.endswith('_contract.html') or fn.endswith('_contract_es.html'):
                contract_files.append(fn)
                # Build the formatted label
                label = fn.replace('_contract_es.html', '').replace('_contract.html', '').replace('_', ' ')
                # Format camelCase words by inserting spaces before uppercase letters
                # and also insert space before "and" when preceded by lowercase letter
                label = re.sub(r'([a-z])([A-Z])', r'\1 \2', label)
                label = re.sub(r'([a-z])(and)', r'\1 \2', label)
                label = label.title()
                contract_choices.append((fn, label))

    context = {
        'case': case,
        'templates': templates,
        'contract_files': contract_files,
        'contract_choices': contract_choices,
        'contract_files_json': json.dumps(contract_files),

    'db_templates': db_templates,
    }
    # Build a small, opinionated placeholder mapping so the frontend can autofill common keys
    placeholder_map = {}
    # Client and attorney
    if client_assignment:
        placeholder_map.update({
            'CLIENT': client_assignment.user.get_full_name(),
            'CLIENT_NAME': client_assignment.user.get_full_name(),
            'CLIENT_EMAIL': client_assignment.user.email or ''
        })
    if attorney_assignment:
        placeholder_map.update({
            'ATTORNEY': attorney_assignment.user.get_full_name(),
            'ATTORNEY_NAME': attorney_assignment.user.get_full_name()
        })

    # Case-related
    placeholder_map.update({
        'CASE_TITLE': case.case_title or '',
        'CASE_DESCRIPTION': case.description or ''
    })

    # Date/time defaults
    now = timezone.now()
    try:
        month_name = now.strftime('%B')
    except Exception:
        month_name = ''
    placeholder_map.update({
        'DAY': str(now.day),
        'MONTH': month_name,
        'YEAR': str(now.year),
        'TIME': now.strftime('%H:%M')
    })

    # Helpful aliases for common roles found in templates
    # Map a few role-like placeholders to the client (best-effort)
    if client_assignment:
        placeholder_map.setdefault('LESSOR', client_assignment.user.get_full_name())
        placeholder_map.setdefault('LESSEE', client_assignment.user.get_full_name())
        placeholder_map.setdefault('SELLER', client_assignment.user.get_full_name())
        placeholder_map.setdefault('BUYERS', client_assignment.user.get_full_name())

    context['placeholder_map_json'] = json.dumps(placeholder_map)
    # Provide case participants for role dropdowns (id + display name)
    participants = []
    for a in case.assignments.select_related('user').all():
        participants.append({'id': a.user.pk, 'name': a.user.get_full_name() or a.user.username})
    context['participants_json'] = json.dumps(participants)
    return render(request, 'cases/generate_document.html', context)

# --- View 12: Signature Request Page (Attorney-facing) ---
@login_required
def signature_request_view(request, doc_pk):
    doc = get_object_or_404(Document, pk=doc_pk)
    case = doc.case

    # --- Security & Permission Check ---
    is_admin = request.user.roles.filter(name='Admin').exists()
    is_assigned = CaseAssignment.objects.filter(case=case, user=request.user).exists()
    
    if not (is_admin or is_assigned):
        messages.error(request, "You do not have permission to access this.")
        return redirect('users:dashboard')

    is_attorney = request.user.roles.filter(name='Attorney').exists()
    if not (is_admin or is_attorney):
        messages.error(request, "Only attorneys can send signature requests.")
        return redirect('cases:case-detail', pk=case.pk)

    # --- Handle Form POST ---
    if request.method == 'POST':
        signer_id = request.POST.get('user_id')
        if not signer_id:
            messages.error(request, "Please select a user to send the request to.")
        else:
            signer = get_object_or_404(User, pk=signer_id)
            
            # Create the SignatureRequest
            sig_request = SignatureRequest.objects.create(
                document=doc,
                signer=signer,
                requested_by=request.user,
                status=SignatureRequest.Status.PENDING
            )
            
            # Build the secure signing link
            sign_link = request.build_absolute_uri(
                reverse('cases:signing-page', kwargs={'token': sig_request.token})
            )
            
            # TODO: Email this link. For now, we'll show it to the attorney.
            messages.success(request, f"Signature request sent to {signer.username}.")
            messages.info(request, f"Signing Link (for demo): {sign_link}")
            return redirect('cases:case-detail', pk=case.pk)

    # --- Handle GET ---
    potential_signers = case.assignments.exclude(user=request.user)
    context = {
        'doc': doc,
        'case': case,
        'potential_signers': potential_signers,
    }
    return render(request, 'cases/signature_request.html', context)


# --- View 13: Signing Page (Client-facing) ---
@login_required
def signing_page_view(request, token):
    # Find the request by its secure token
    sig_request = get_object_or_404(SignatureRequest, token=token)
    doc = sig_request.document

    # --- Security & Validation Checks ---
    # 1. Check if user is the correct signer
    if request.user != sig_request.signer:
        messages.error(request, "You are not authorized to sign this document.")
        return redirect('users:dashboard')
    
    # 2. Check if it's still pending
    if sig_request.status != SignatureRequest.Status.PENDING:
        messages.warning(request, "This document has already been signed or the request was cancelled.")
        return redirect('cases:case-detail', pk=doc.case.pk)

    # --- Handle the POST (signing) action ---
    if request.method == 'POST':
        # Check if the "agree" box was ticked
        if 'agree' not in request.POST:
            messages.error(request, "You must agree to the terms to sign.")
        else:
            # --- SUCCESS ---
            # 1. Update the request
            sig_request.status = SignatureRequest.Status.SIGNED
            sig_request.signed_at = timezone.now()
            sig_request.save()
            
            # 2. Create the permanent audit log
            DocumentLog.objects.create(
                document=doc,
                user=request.user,
                action="Signed",
                details=f"Signed via e-signature workflow."
            )
            
            messages.success(request, "Document successfully signed.")
            return redirect('cases:case-detail', pk=doc.case.pk)

    # --- Handle GET ---
    context = {
        'doc': doc,
        'sig_request': sig_request
    }
    return render(request, 'cases/signing_page.html', context)

# --- View 14: Calendar ---
@login_required
def calendar_events_api(request):
    """
    Returns calendar events as JSON for the logged-in user.
    FullCalendar will request this and display the events.
    """
    user = request.user
    events = []
    
    # Get all meetings the user is involved in (either as organizer or participant)
    meetings = Meeting.objects.filter(
        participants=user
    ).select_related('case', 'organizer').values('id', 'title', 'scheduled_time', 'duration_minutes', 'meeting_type', 
        'description', 'case__case_title', 'case__pk', 'organizer__username', 
        'organizer__first_name', 'organizer__last_name')
    
    for meeting in meetings:
        end_time = meeting['scheduled_time'] + timezone.timedelta(minutes=meeting['duration_minutes'])

        # Get participant names for this meeting
        meeting_obj = Meeting.objects.get(id=meeting['id'])
        participants = [p.get_full_name() or p.username for p in meeting_obj.participants.all()]
        
        # Format organizer name
        organizer_name = meeting['organizer__first_name'] or ''
        if meeting['organizer__last_name']:
            organizer_name = f"{organizer_name} {meeting['organizer__last_name']}".strip()
        if not organizer_name:
            organizer_name = meeting['organizer__username']

        events.append({
            'id': f'meeting-{meeting["id"]}',
            'title': f"📞 {meeting['title']}",
            'start': meeting['scheduled_time'].isoformat(),
            'end': end_time.isoformat(),
            'type': 'meeting',
            'backgroundColor': '#c4a24c',  # Your owl gold color
            'borderColor': '#c4a24c',
            'extendedProps': {
                'meetingId': meeting['id'],
                'caseTitle': meeting['case__case_title'],
                'caseId': meeting['case__pk'],
                'meetingType': meeting['meeting_type'],
                'description': meeting['description'] or 'No description provided',
                'duration': meeting['duration_minutes'],
                'organizer': organizer_name,
                'participants': ', '.join(participants),
                'scheduledTime': meeting['scheduled_time'].strftime('%B %d, %Y at %I:%M %p'),
            }
        })
    
    # Get all document due dates assigned to the user or linked to their cases
    documents = DocumentDueDate.objects.filter(
        case__assignments__user=user
    ).select_related('case').values('id', 'document_name', 'due_date', 'is_completed', 'case__case_title').distinct()
    
    for doc in documents:
        events.append({
            'id': f'document-{doc["id"]}',
            'title': f"📄 {doc['document_name']}",
            'start': doc['due_date'].isoformat(),
            'type': 'document',
            'backgroundColor': '#ff6b6b' if not doc['is_completed'] else '#51cf66',
            'borderColor': '#ff6b6b' if not doc['is_completed'] else '#51cf66',
            'extendedProps': {
                'caseTitle': doc['case__case_title'],
            }
        })
    
    return JsonResponse(events, safe=False)

# --- View 15: Meetings ---
@login_required
# @user_passes_test(is_admin) //we want admin and attorneys to create meetings
def create_meeting_view(request, case_pk=None):
    """
    Create a meeting for a case.
    
    If case_pk is provided (from case detail), that case is pre-selected.
    If case_pk is NOT provided (from dashboard), admin must select a case.
    
    This allows the same view to work from both places.
    """

    # --- Permission Check: Only Admins and Attorneys can create meetings ---
    is_admin = request.user.roles.filter(name='Admin').exists()
    is_attorney = request.user.roles.filter(name='Attorney').exists()
    
    if not (is_admin or is_attorney):
        messages.error(request, "Only admins and attorneys can schedule meetings.")
        return redirect('users:dashboard')
    
    # Pre-select a case if we came from case detail
    case = None
    if case_pk:
        case = get_object_or_404(Case, pk=case_pk)
    
    if request.method == 'POST':
        form = MeetingForm(request.POST)
        if form.is_valid():
            meeting = form.save(commit=False)
            meeting.organizer = request.user
            meeting.save()
            # Save the many-to-many participants
            form.save_m2m()
            
            messages.success(request, f"Meeting '{meeting.title}' created successfully!")
            
            # Redirect back to the case detail (or dashboard if no case)
            if case:
                return redirect('cases:case-detail', pk=case.pk)
            else:
                return redirect('cases:case-dashboard')
    else:
        form = MeetingForm()
        
        # If we have a pre-selected case, set it and disable the dropdown
        if case:
            form.fields['case'].initial = case
            form.fields['case'].widget = forms.HiddenInput()
            form.fields['case'].help_text = "This case was selected from the case detail page."
            
            # Pre-filter participants to show only users assigned to THIS case
            form.fields['participants'].queryset = User.objects.filter(
                case_assignments__case=case
            ).distinct()
        else:
            # No case selected yet, show all cases
            form.fields['case'].help_text = "Select the case for this meeting."
    
    context = {
        'form': form,
        'case': case,
    }
    return render(request, 'cases/create_meeting.html', context)

def calendar_view(request):
    return render(request, 'cases/calendar.html', {})

# --- List and Create Templates (GET /api/templates/, POST /api/templates/) ---
class ContractTemplateListCreateView(generics.ListCreateAPIView):
    queryset = ContractTemplate.objects.all().order_by('-updated_at')
    serializer_class = ContractTemplateSerializer
    permission_classes = [IsAuthenticated] # Ensures only logged-in users can access

    def perform_create(self, serializer):
        # Set the 'created_by' field automatically to the logged-in user
        serializer.save(created_by=self.request.user)
    
    # Optionally filter to only show public templates and the user's own templates
    def get_queryset(self):
        user = self.request.user

        # 1. Allow staff/superusers to see ALL templates
        if user.is_authenticated and (user.is_staff or user.is_superuser):
            return ContractTemplate.objects.all().order_by('-updated_at')

        # 2. Default filtering for standard users (Public OR User's own)
        if user.is_authenticated:
            return ContractTemplate.objects.filter(
                models.Q(is_public=True) | models.Q(created_by=user)
            ).order_by('-updated_at')
        
        # 3. For anonymous users (only public templates)
        return ContractTemplate.objects.filter(is_public=True).order_by('-updated_at')
    
# --- Retrieve, Update, and Destroy a Template (GET/PUT/PATCH/DELETE /api/templates/{id}/) ---
class ContractTemplateRetrieveUpdateDestroyView(generics.RetrieveUpdateDestroyAPIView):
    queryset = ContractTemplate.objects.all()
    serializer_class = ContractTemplateSerializer
    permission_classes = [IsAuthenticated]
    
    # NOTE: You may want to add custom permission to ensure only the creator or an admin can edit/delete
    # For example, using a custom permission class like: IsOwnerOrReadOnly

def template_list(request):
    # This view just renders the HTML page. 
    # The data fetching is done by JavaScript via the API.
    return render(request, 'cases/template_list.html', {})



class ContractTemplateDownloadView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, pk):
        # 1. Get the template
        template = get_object_or_404(ContractTemplate, pk=pk)
        
        # 2. Define an empty context (since this is a blank template download)
        context = {}

        # 3. Generate the PDF using your helper function
        # (Ensure generate_document_from_template is imported!)
        file_buffer, filename = generate_document_from_template(
            template.content, 
            template.name,
            context_data=context
        )

        # 4. Return the file as a PDF attachment
        response = FileResponse(file_buffer, as_attachment=True)
        response['Content-Type'] = 'application/pdf'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'

        return response
    
# Add to bottom of cases/views.py
@login_required
def get_db_template_content(request, pk):
    """Returns the raw HTML content of a database ContractTemplate"""
    template = get_object_or_404(ContractTemplate, pk=pk)
    return HttpResponse(template.content, content_type='text/html')

@login_required
def client_list_view(request):
    User = get_user_model()
    # Filter for users who have the 'Client' role
    users = User.objects.filter(roles__name='Client').distinct().order_by('last_name')
    context = {
        'users': users,
        'page_title': 'Clients',
        'icon': 'fas fa-user-friends'
    }
    return render(request, 'cases/user_list.html', context)

@login_required
def attorney_list_view(request):
    User = get_user_model()
    # Filter for users who have the 'Attorney' role
    users = User.objects.filter(roles__name='Attorney').distinct().order_by('last_name')
    context = {
        'users': users,
        'page_title': 'Attorneys',
        'icon': 'fas fa-gavel'
    }
    return render(request, 'cases/user_list.html', context)

@login_required
def user_profile_view(request, pk):
    User = get_user_model()
    profile_user = get_object_or_404(User, pk=pk)
    
    # 1. Get all cases this user is assigned to
    assignments = CaseAssignment.objects.filter(user=profile_user).select_related('case')
    
    # 2. Determine "Related People" logic
    related_people = []
    related_title = ""
    
    is_attorney_profile = profile_user.roles.filter(name='Attorney').exists()
    is_client_profile = profile_user.roles.filter(name='Client').exists()

    # Get the IDs of cases this user is part of
    case_ids = assignments.values_list('case_id', flat=True)

    if is_attorney_profile:
        # If looking at an Attorney, show their Clients
        related_title = "Clients"
        related_people = User.objects.filter(
            case_assignments__case__id__in=case_ids,
            roles__name='Client'
        ).distinct()
        
    elif is_client_profile:
        # If looking at a Client, show their Attorneys
        related_title = "Attorneys"
        related_people = User.objects.filter(
            case_assignments__case__id__in=case_ids,
            roles__name='Attorney'
        ).distinct()

    context = {
        'profile_user': profile_user,
        'assignments': assignments,
        'related_people': related_people,
        'related_title': related_title,
    }
    return render(request, 'cases/user_profile.html', context)

@login_required
# @user_passes_test(is_admin) # Ensure you have your admin check here
def user_management_view(request):
    User = get_user_model()
    
    # 1. Start with all users
    users = User.objects.all().order_by('-date_joined')
    
    # 2. Search Filter (checks username, first name, last name, email)
    search_query = request.GET.get('q', '')
    if search_query:
        users = users.filter(
            Q(username__icontains=search_query) |
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(email__icontains=search_query)
        )

    # 3. Role Filter (checks the Role model relation)
    role_filter = request.GET.get('role', '')
    if role_filter:
        users = users.filter(roles__name=role_filter)
    
    context = {
        'users': users,
        'search_query': search_query,
        'role_filter': role_filter
    }
    return render(request, 'cases/user_management_list.html', context)

# --- 1. Public Consultation Request View ---
def request_consultation_view(request):
    initial_data = {}
    
    # If accessed via ?attorney_id=5, pre-select that attorney
    attorney_id = request.GET.get('attorney_id')
    if attorney_id:
        initial_data['attorney'] = attorney_id

    if request.method == 'POST':
        form = ConsultationForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Your consultation request has been sent. We will contact you shortly.")
            return redirect('homepage')
    else:
        form = ConsultationForm(initial=initial_data)

    return render(request, 'cases/consultation_form.html', {'form': form})

# --- 2. Attorney Actions ---
@login_required
def update_consultation_status(request, pk, status):
    consultation = get_object_or_404(ConsultationRequest, pk=pk)
    
    # Security check: Ensure current user is the assigned attorney or an admin
    if request.user != consultation.attorney and not request.user.is_superuser:
        messages.error(request, "You cannot manage this request.")
        return redirect('cases:case-dashboard')

    consultation.status = status
    consultation.save()
    messages.success(request, f"Request marked as {status}.")
    return redirect('cases:case-dashboard')

# --- 2. Schedule & Email Action ---
@login_required
def schedule_consultation_action(request, pk):
    consultation = get_object_or_404(ConsultationRequest, pk=pk)
    
    if request.method == 'POST':
        form = ConsultationScheduleForm(request.POST)
        if form.is_valid():
            # Extract data
            m_type = form.cleaned_data['meeting_type']
            time = form.cleaned_data['scheduled_time']
            link = form.cleaned_data['booking_link']
            msg = form.cleaned_data['additional_message']
            
            # Construct Email Content
            subject = f"Update on your Consultation Request: {consultation.service_needed[:30]}..."
            
            email_body = f"Dear {consultation.name},\n\n"
            email_body += f"Attorney {request.user.get_full_name()} has reviewed your request.\n\n"
            email_body += f"Meeting Type: {m_type}\n"
            
            if time:
                email_body += f"Confirmed Time: {time.strftime('%B %d, %Y at %I:%M %p')}\n"
                consultation.status = 'Scheduled' # Auto-update status
            
            if link:
                email_body += f"Please schedule your time here: {link}\n"
            
            if msg:
                email_body += f"\nMessage from Attorney:\n{msg}\n"
            
            email_body += "\nBest regards,\nThe Owlery Legal Team"

            # Send Email (Print to console if email backend not set up)
            try:
                if consultation.email:
                    send_mail(
                        subject,
                        email_body,
                        settings.DEFAULT_FROM_EMAIL,
                        [consultation.email],
                        fail_silently=False,
                    )
                    messages.success(request, f"Invitation sent to {consultation.email}")
                else:
                    messages.warning(request, "Client has no email address.")
            except Exception as e:
                messages.error(request, f"Error sending email: {e}")

            consultation.save()
            return redirect('cases:consultation-detail', pk=consultation.pk)
    
    return redirect('cases:consultation-detail', pk=consultation.pk)

# --- View 21: Consultation Details & Scheduling ---
@login_required
def consultation_detail_view(request, pk):
    consultation = get_object_or_404(ConsultationRequest, pk=pk)
    
    # Security check
    if request.user != consultation.attorney and not request.user.is_superuser:
        messages.error(request, "Access denied.")
        return redirect('cases:case-dashboard')

    schedule_form = ConsultationScheduleForm()
    return render(request, 'cases/consultation_detail.html', {'consultation': consultation, 'schedule_form': schedule_form})

@login_required
def case_directory_view(request):
    # 1. Check Permissions
    is_admin = request.user.roles.filter(name='Admin').exists()
    is_attorney = request.user.roles.filter(name='Attorney').exists()
    
    if not (is_admin or is_attorney):
        messages.error(request, "Access denied.")
        return redirect('users:dashboard')

    # 2. Fetch ALL Cases
    cases = Case.objects.all().select_related('current_stage').prefetch_related('assignments__user').order_by('-date_filed')
    
    # 3. Filter for Attorneys (See only assigned cases)
    if not is_admin:
        cases = cases.filter(assignments__user=request.user)

    # --- 4. NEW: Search Functionality ---
    search_query = request.GET.get('q', '')
    if search_query:
        # Create a query that looks in Title or Assigned User Names
        query = Q(case_title__icontains=search_query) | \
                Q(assignments__user__first_name__icontains=search_query) | \
                Q(assignments__user__last_name__icontains=search_query) | \
                Q(assignments__user__username__icontains=search_query)
        
        # Add special handling for Status keywords
        if search_query.lower() == 'active':
            query |= Q(is_archived=False)
        elif search_query.lower() == 'archived':
            query |= Q(is_archived=True)
            
        cases = cases.filter(query).distinct()

    context = {
        'cases': cases,
        'search_query': search_query, # Pass this back to keep the text in the box
    }
    return render(request, 'cases/case_directory.html', context)